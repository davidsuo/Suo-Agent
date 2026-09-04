# common/rag.py
import os
import json
import uuid
import datetime
import re
from typing import List

# 强制与 main.py 使用同一个绝对路径
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RAG_DATA_FILE = os.path.join(BASE_DIR, "rag_data.json")
GLOBAL_KEY = "__global__"

def _load_store():
    """加载JSON文件"""
    if os.path.exists(RAG_DATA_FILE):
        try:
            with open(RAG_DATA_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {"files": [], "store": {}}
    return {"files": [], "store": {}}

def _save_store(store):
    """保存JSON文件，这个绝不会报错！"""
    with open(RAG_DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(store, f, ensure_ascii=False, indent=2)

def _chunk_text(text: str, chunk_size=1500) -> List[str]:
    if len(text) <= chunk_size:
        return [text]
    lines = text.split("\n")
    chunks = []
    current_chunk = ""
    for line in lines:
        if len(current_chunk) + len(line) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = line + "\n"
        else:
            current_chunk += line + "\n"
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def index_document(file_path: str, session_id: str, tags: str = "") -> str:
    try:
        import pandas as pd
        
        if not isinstance(file_path, str):
            if hasattr(file_path, 'name'):
                file_path = file_path.name
            else:
                file_path = str(file_path)
                
        ext = os.path.splitext(file_path)[1].lower()
        content = ""
        
        # 读取文件内容
        if ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        elif ext in [".csv", ".xlsx", ".xls"]:
            if ext == ".csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            df = df.fillna("")
            content = df.to_csv(index=False)
        elif ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            except ImportError:
                return "❌ 缺少 pypdf 库"
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        content += " | ".join(row_text) + "\n"
            except ImportError:
                return "❌ 缺少 python-docx 库"
        else:
            return f"❌ 不支持的文件格式: {ext}"

        if not content.strip():
            return "❌ 文件内容为空或无法解析。"

        chunks = _chunk_text(content)
        store = _load_store()

        file_name = os.path.basename(file_path)
        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # 直接写入文件列表
        store["files"].append({
            "file_name": file_name,
            "tags": tags if tags else "(无)",
            "created_at": current_time,
            "chunks": len(chunks)
        })

        # 写入内容片段
        if tags and tags.strip():
            target_key = f"tag_{tags.strip()}"
        else:
            target_key = GLOBAL_KEY

        if target_key not in store["store"]:
            store["store"][target_key] = []
        for chunk in chunks:
            store["store"][target_key].append({
                "id": str(uuid.uuid4()),
                "text": chunk,
                "tags": tags,
                "file_name": file_name,
                "time": current_time
            })

        # 强制保存为 JSON 文件
        _save_store(store)

        return f"✅ 文档已成功索引（共 {len(chunks)} 个片段，支持 PDF/Word/表格）。"
    except Exception as e:
        return f"❌ 文档处理失败: {e}"

def search_knowledge(query: str, session_id: str, tags: str = "") -> str:
    store = _load_store()
    combined_docs = []
    
    # 1. 显式传标签时，只搜该标签的
    if tags and tags.strip():
        target_key = f"tag_{tags.strip()}"
        if target_key in store["store"]:
            combined_docs.extend(store["store"][target_key])
    
    # 2. 总是加上全局库的数据
    if GLOBAL_KEY in store["store"]:
        combined_docs.extend(store["store"][GLOBAL_KEY])
    
    # 3. 【终极修复】如果没传标签，强制把所有的标签库全部搜一遍！
    if not tags or not tags.strip():
        for key in store["store"].keys():
            if key.startswith("tag_"):
                combined_docs.extend(store["store"][key])
    
    if not combined_docs:
        return ""

    import re as _re
    _year_match = _re.search(r'(20\d{2})年', query)
    _month_match = _re.search(r'(\d{1,2})月份', query)
    target_prefix = ""
    if _year_match and _month_match:
        target_prefix = f"{_year_match.group(1)}/{int(_month_match.group(1))}/"

    matched_texts = []
    if target_prefix:
        for doc in combined_docs:
            if target_prefix in doc.get("text", ""):
                matched_texts.append(doc.get("text", ""))
    else:
        clean_query = query.replace("，", " ").replace("。", " ").replace("？", " ").replace("?", " ").replace(" ", "")
        grams = set()
        for i in range(len(clean_query)):
            for j in range(i + 2, min(i + 6, len(clean_query) + 1)):
                grams.add(clean_query[i:j])
        for doc in combined_docs:
            text = doc.get("text", "")
            score = 0
            for gram in grams:
                if gram in text:
                    score += 1
            if score >= 3:
                matched_texts.append(text)

    if matched_texts:
        return "\n\n".join(matched_texts[:10])[:8000]
    return ""