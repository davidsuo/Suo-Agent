# common/rag_v2.py
import os
import json
import uuid
import datetime
from typing import List

try:
    import chromadb
    from chromadb.config import Settings
except ImportError:
    print("❌ 请先安装依赖: pip install chromadb")

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHROMA_DIR = os.path.join(BASE_DIR, "chroma_db")
RAG_DATA_FILE = os.path.join(BASE_DIR, "rag_data.json")

# 初始化 ChromaDB 客户端（持久化模式）
_client = chromadb.PersistentClient(path=CHROMA_DIR)
_collection = _client.get_or_create_collection(name="enterprise_kb")

# 简单的文档感知分块
def smart_chunk_text(text: str, max_chunk_size: int = 1500) -> List[str]:
    if len(text) <= max_chunk_size:
        return [text]
    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) > max_chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = para + "\n"
        else:
            current_chunk += para + "\n"
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def index_document_v2(file_path: str, tags: str = ""):
    """新的索引入库逻辑：读取 -> 智能分块 -> 向量化入库 -> 清理旧数据 -> 同步Json列表与内容"""
    try:
        import pandas as pd
        ext = os.path.splitext(file_path)[1].lower()
        content = ""
        if ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        elif ext in [".csv", ".xlsx", ".xls"]:
            if ext == ".csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            df = df.fillna("")
            content = df.to_csv(index=False)
        elif ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            except ImportError:
                return "❌ 缺少 pypdf 库"
        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document(file_path)
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        content += " | ".join(row_text) + "\n"
            except ImportError:
                return "❌ 缺少 python-docx 库"
        else:
            return f"❌ 不支持的文件格式: {ext}"

        if not content.strip():
            return "❌ 文件内容为空。"

        chunks = smart_chunk_text(content)
        file_name = os.path.basename(file_path)
        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # ================= 【新增：向量刷新核心逻辑】 =================
        # 1. 删除 ChromaDB 中同名旧向量（防止数据堆积损坏 HNSW 索引）
        try:
            _collection.delete(where={"file_name": file_name})
        except Exception as e:
            print(f"###DEBUG### 清理旧向量失败: {e}")

        # 2. 清理 rag_data.json 中的旧记录
        store = {"files": [], "store": {}}
        if os.path.exists(RAG_DATA_FILE):
            try:
                with open(RAG_DATA_FILE, "r", encoding="utf-8") as f:
                    store = json.load(f)
            except Exception:
                store = {"files": [], "store": {}}
        
        # 从 files 列表中移除旧文件
        store["files"] = [f for f in store.get("files", []) if f.get("file_name") != file_name]
        
        # 从 store 字典中移除旧文件的所有内容
        for key in list(store.get("store", {}).keys()):
            store["store"][key] = [doc for doc in store["store"][key] if doc.get("file_name") != file_name]

        # 3. 写入全新的索引和内容
        ids = [str(uuid.uuid4()) for _ in chunks]
        metadatas = [{"file_name": file_name, "tags": tags, "time": current_time} for _ in chunks]
        _collection.add(
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )

        # 4. 同步更新 rag_data.json
        store.setdefault("files", []).append({
            "file_name": file_name,
            "tags": tags if tags else "(无)",
            "created_at": current_time,
            "chunks": len(chunks)
        })

        target_key = f"tag_{tags.strip()}" if tags and tags.strip() else "__global__"
        if target_key not in store["store"]:
            store["store"][target_key] = []
        for chunk in chunks:
            store["store"][target_key].append({
                "id": str(uuid.uuid4()),
                "text": chunk,
                "tags": tags,
                "file_name": file_name,
                "time": current_time
            })

        with open(RAG_DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(store, f, ensure_ascii=False, indent=2)

        return f"✅ [V2] 文档已成功入库到ChromaDB（共 {len(chunks)} 个智能分块）。"
    except Exception as e:
        return f"❌ [V2] 文档处理失败: {e}"

def search_knowledge_v2(query: str, tags: str = ""):
    """【V2终极版】直接读取同步后的JSON数据进行混合检索（极速，规避ChromaDB拉全量慢+where=None坑）"""
    try:
        import re as _re
        # 读取同步后的 JSON 数据
        store = {"files": [], "store": {}}
        if os.path.exists(RAG_DATA_FILE):
            with open(RAG_DATA_FILE, "r", encoding="utf-8") as f:
                store = json.load(f)

        combined_docs = []
        # 如果传了标签，仅搜索对应标签下的内容
        if tags and tags.strip():
            target_key = f"tag_{tags.strip()}"
            if target_key in store["store"]:
                combined_docs.extend(store["store"][target_key])
        # 即使没传标签，也加上全局数据
        if "__global__" in store["store"]:
            combined_docs.extend(store["store"]["__global__"])

        # 如果没传标签，则搜索所有标签下的内容（等同于V1的终极修复逻辑）
        if not tags or not tags.strip():
            for key in store["store"].keys():
                if key.startswith("tag_"):
                    combined_docs.extend(store["store"][key])
            if "__global__" in store["store"]:
                combined_docs.extend(store["store"]["__global__"])

        if not combined_docs:
            return ""

        # 解析查询中的年份和月份
        _year_match = _re.search(r'(20\d{2})', query)
        _month_match = _re.search(r'(\d{1,2})月份', query)
        target_prefix = ""
        if _year_match and _month_match:
            target_prefix = f"{_year_match.group(1)}/{int(_month_match.group(1))}/"

        matched_texts = []
        # 精准的日期/数值匹配模式
        if target_prefix:
            for doc in combined_docs:
                if target_prefix in doc.get("text", ""):
                    matched_texts.append(doc.get("text", ""))
        # 语义匹配模式
        else:
            clean_query = query.replace("，", " ").replace("。", " ").replace("？", " ").replace("?", " ").replace(" ", "")
            grams = set()
            for i in range(len(clean_query)):
                for j in range(i + 2, min(i + 6, len(clean_query) + 1)):
                    grams.add(clean_query[i:j])
            for doc in combined_docs:
                text = doc.get("text", "")
                score = 0
                for gram in grams:
                    if gram in text:
                        score += 1
                if score >= 3:
                    matched_texts.append(text)

        if matched_texts:
            return "\n\n".join(matched_texts[:10])[:20000]
        return ""
    except Exception as e:
        print(f"###DEBUG### V2混合检索失败: {e}")
        return ""